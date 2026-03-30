# =====================================================
# ingestion/parser.py
# Extracts plain text from PDF, DOCX, and TXT files.
#
# WHY separate this into its own file?
# Each file type needs a different library.
# Keeping them isolated means adding a new format
# (e.g. XLSX) only requires adding one function here.
# =====================================================

import io
from loguru import logger


def parse_pdf(file_bytes: bytes) -> str:
    """
    Extract text from a PDF file.

    HOW IT WORKS:
    PDFs store text in a complex binary format.
    pypdf reads each page and extracts the text layer.
    Note: scanned PDFs (images) won't work — they need OCR.
    For this project we assume text-based PDFs.
    """
    try:
        from pypdf import PdfReader

        # Load PDF from bytes (we get bytes from Supabase Storage)
        reader = PdfReader(io.BytesIO(file_bytes))

        text_parts = []
        for page_num, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text and page_text.strip():
                # Add page marker — useful for citations later
                text_parts.append(f"[Page {page_num + 1}]\n{page_text}")

        full_text = "\n\n".join(text_parts)

        if not full_text.strip():
            raise ValueError(
                "No text extracted from PDF. "
                "It may be a scanned image — text-based PDFs only."
            )

        logger.info(f"PDF parsed | pages={len(reader.pages)} chars={len(full_text)}")
        return full_text

    except Exception as e:
        logger.error(f"PDF parsing failed: {e}")
        raise


def parse_docx(file_bytes: bytes) -> str:
    """
    Extract text from a DOCX (Word) file.

    HOW IT WORKS:
    DOCX files are ZIP archives containing XML.
    python-docx reads paragraphs and tables from the XML.
    """
    try:
        from docx import Document

        doc = Document(io.BytesIO(file_bytes))
        text_parts = []

        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)

        # Extract text from tables too
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip()
                    for cell in row.cells
                    if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        full_text = "\n\n".join(text_parts)

        if not full_text.strip():
            raise ValueError("No text found in DOCX file.")

        logger.info(f"DOCX parsed | paragraphs={len(doc.paragraphs)} chars={len(full_text)}")
        return full_text

    except Exception as e:
        logger.error(f"DOCX parsing failed: {e}")
        raise


def parse_txt(file_bytes: bytes) -> str:
    """
    Read plain text file.
    Tries UTF-8 first, falls back to latin-1.
    """
    try:
        try:
            text = file_bytes.decode("utf-8")
        except UnicodeDecodeError:
            text = file_bytes.decode("latin-1")

        if not text.strip():
            raise ValueError("TXT file is empty.")

        logger.info(f"TXT parsed | chars={len(text)}")
        return text

    except Exception as e:
        logger.error(f"TXT parsing failed: {e}")
        raise


def parse_document(file_bytes: bytes, file_type: str) -> str:
    """
    Main entry point — routes to the correct parser
    based on file extension.

    Args:
        file_bytes: Raw file content from storage
        file_type: 'pdf', 'docx', or 'txt'

    Returns:
        Extracted plain text string
    """
    parsers = {
        "pdf":  parse_pdf,
        "docx": parse_docx,
        "txt":  parse_txt,
    }

    parser = parsers.get(file_type.lower())

    if not parser:
        raise ValueError(
            f"Unsupported file type: '{file_type}'. "
            f"Supported: {list(parsers.keys())}"
        )

    return parser(file_bytes)